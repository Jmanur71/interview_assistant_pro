"""Resume file parser supporting multiple formats: PDF, DOCX, DOC, TXT, RTF"""

import os
import re
from typing import Tuple, Optional


class ResumeParser:
    """Parse resume files in various formats and extract text content."""

    SUPPORTED_FORMATS = {
        ".pdf": "PDF Document",
        ".docx": "Word Document (2007+)",
        ".doc": "Word Document (97-2003)",
        ".txt": "Plain Text",
        ".rtf": "Rich Text Format",
    }

    @staticmethod
    def get_supported_formats_filter() -> str:
        """Return file dialog filter string for supported formats."""
        filters = []
        for ext, desc in ResumeParser.SUPPORTED_FORMATS.items():
            filters.append(f"{desc} (*{ext})")
        filters.append("All files (*.*)")
        return ";;".join(filters)

    @staticmethod
    def parse(file_path: str) -> Tuple[Optional[str], Optional[str]]:
        """
        Parse resume file and return (content, error_message).
        Returns (content, None) on success, (None, error_msg) on failure.
        """
        try:
            ext = os.path.splitext(file_path)[1].lower()

            if ext == ".pdf":
                return ResumeParser._parse_pdf(file_path)
            elif ext == ".docx":
                return ResumeParser._parse_docx(file_path)
            elif ext == ".doc":
                return ResumeParser._parse_doc(file_path)
            elif ext == ".txt":
                return ResumeParser._parse_txt(file_path)
            elif ext == ".rtf":
                return ResumeParser._parse_rtf(file_path)
            else:
                return None, f"Unsupported file format: {ext}"

        except Exception as e:
            return None, f"Error parsing file: {str(e)}"

    @staticmethod
    def _parse_pdf(file_path: str) -> Tuple[Optional[str], Optional[str]]:
        """Extract text from PDF files."""
        try:
            from pypdf import PdfReader

            reader = PdfReader(file_path)
            text = ""
            for page_num, page in enumerate(reader.pages):
                try:
                    text += f"\n--- Page {page_num + 1} ---\n"
                    page_text = page.extract_text() or ""
                    text += page_text
                except Exception as e:
                    text += f"[Error extracting page {page_num + 1}: {e}]\n"

            if not text.strip():
                return None, "PDF file is empty or contains no extractable text"

            return text.strip(), None

        except ImportError:
            return None, "PyPDF library not installed. Run: pip install pypdf"
        except Exception as e:
            return None, f"Error reading PDF: {str(e)}"

    @staticmethod
    def _parse_docx(file_path: str) -> Tuple[Optional[str], Optional[str]]:
        """Extract text from DOCX files (Word 2007+)."""
        try:
            from docx import Document

            doc = Document(file_path)
            text = ""

            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text += para.text + "\n"

            # Extract tables
            if doc.tables:
                text += "\n--- Tables ---\n"
                for table in doc.tables:
                    for row in table.rows:
                        cells = [cell.text.strip() for cell in row.cells]
                        text += " | ".join(cells) + "\n"

            if not text.strip():
                return None, "DOCX file is empty or contains no text"

            return text.strip(), None

        except ImportError:
            return None, "python-docx library not installed. Run: pip install python-docx"
        except Exception as e:
            return None, f"Error reading DOCX: {str(e)}"

    @staticmethod
    def _parse_doc(file_path: str) -> Tuple[Optional[str], Optional[str]]:
        """Extract text from DOC files (Word 97-2003)."""
        try:
            # Try using python-docx first (some .doc files are compatible)
            from docx import Document

            try:
                doc = Document(file_path)
                text = ""
                for para in doc.paragraphs:
                    if para.text.strip():
                        text += para.text + "\n"
                if text.strip():
                    return text.strip(), None
            except Exception:
                pass

            # Fallback: Try reading as binary and extract text
            with open(file_path, "rb") as f:
                content = f.read()
                # Extract printable ASCII text (crude but works for basic .doc files)
                text = re.sub(rb"[^\x20-\x7E\n\r\t]", b" ", content).decode("utf-8", errors="ignore")
                text = re.sub(r"\s+", " ", text).strip()

                if len(text) > 20:
                    return text, None
                else:
                    return None, "DOC file could not be parsed. Try converting to DOCX or PDF."

        except ImportError:
            # Fallback without python-docx
            try:
                with open(file_path, "rb") as f:
                    content = f.read()
                    text = re.sub(rb"[^\x20-\x7E\n\r\t]", b" ", content).decode("utf-8", errors="ignore")
                    text = re.sub(r"\s+", " ", text).strip()
                    if len(text) > 20:
                        return text, None
                    else:
                        return None, "Could not extract text from DOC file"
            except Exception as e:
                return None, f"Error reading DOC: {str(e)}"

    @staticmethod
    def _parse_txt(file_path: str) -> Tuple[Optional[str], Optional[str]]:
        """Extract text from TXT files."""
        try:
            with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                text = f.read().strip()

            if not text:
                return None, "TXT file is empty"

            return text, None

        except Exception as e:
            return None, f"Error reading TXT: {str(e)}"

    @staticmethod
    def _parse_rtf(file_path: str) -> Tuple[Optional[str], Optional[str]]:
        """Extract text from RTF files."""
        try:
            with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                content = f.read()

            # Remove RTF control sequences
            text = re.sub(r"\\['\"][\da-f]{2}", "", content)  # hex escapes
            text = re.sub(r"[\\{}]", " ", text)  # RTF commands
            text = re.sub(r"\s+", " ", text).strip()  # normalize whitespace

            if not text:
                return None, "RTF file is empty or unreadable"

            return text, None

        except Exception as e:
            return None, f"Error reading RTF: {str(e)}"

    @staticmethod
    def get_file_size_info(file_path: str) -> str:
        """Return human-readable file size."""
        try:
            size_bytes = os.path.getsize(file_path)
            for unit in ["B", "KB", "MB"]:
                if size_bytes < 1024:
                    return f"{size_bytes:.1f} {unit}"
                size_bytes /= 1024
            return f"{size_bytes:.1f} GB"
        except Exception:
            return "Unknown size"
